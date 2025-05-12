import os
import io
import tempfile
import asyncio
import uuid
import json
import pandas as pd
import zipfile
import xlsxwriter
import openpyxl
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

from application.dto import CreateDocumentDTO, CreateTemplateDTO, TemplateDataDTO, MergeDocumentsDTO, BatchProcessingDTO
from domain.models import ExcelDocumentInfo, ExcelTemplateInfo, BatchProcessingInfo, MergeInfo
from domain.exceptions import DocumentNotFoundException, TemplateNotFoundException, StorageException
from domain.exceptions import ConversionException, TemplateApplicationException, MergeException
from infrastructure.repository import ExcelDocumentRepository, ExcelTemplateRepository, BatchProcessingRepository, \
    MergeRepository
from infrastructure.minio_client import MinioClient
from infrastructure.rabbitmq_client import RabbitMQClient
from core.config import settings

from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter


class ExcelDocumentService:
    """
    Service xử lý tài liệu Excel.
    """

    def __init__(
            self,
            document_repository: ExcelDocumentRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            document_repository: Repository để làm việc với tài liệu
            minio_client: Client MinIO để lưu trữ tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.document_repository = document_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client
        self.merge_repository = MergeRepository()

    async def create_document(self, dto: CreateDocumentDTO, content: bytes) -> ExcelDocumentInfo:
        """
        Tạo tài liệu mới.

        Args:
            dto: DTO chứa thông tin tài liệu
            content: Nội dung tài liệu

        Returns:
            Thông tin tài liệu đã tạo
        """
        sheet_names = await self._get_sheet_names(content)

        document_info = ExcelDocumentInfo(
            title=dto.title,
            description=dto.description,
            original_filename=dto.original_filename,
            file_size=len(content),
            file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" if dto.original_filename.endswith(
                ".xlsx") else "application/vnd.ms-excel",
            storage_path="",  
            metadata=dto.metadata,
            sheet_names=sheet_names
        )

        document_info = await self.document_repository.save(document_info, content)

        return document_info

    async def _get_sheet_names(self, content: bytes) -> List[str]:
        """
        Lấy danh sách tên sheet từ file Excel.

        Args:
            content: Nội dung file Excel

        Returns:
            Danh sách tên sheet
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                wb = load_workbook(temp_file_path, read_only=True)
                sheet_names = wb.sheetnames
                wb.close()

                os.unlink(temp_file_path)

                return sheet_names
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                print(f"Lỗi khi đọc tên sheet: {str(e)}")
                return []
        except Exception as e:
            print(f"Lỗi khi tạo file tạm: {str(e)}")
            return []

    async def get_documents(self, skip: int = 0, limit: int = 10, search: Optional[str] = None) -> List[
        ExcelDocumentInfo]:
        """
        Lấy danh sách tài liệu.

        Args:
            skip: Số tài liệu bỏ qua
            limit: Số tài liệu tối đa trả về
            search: Từ khóa tìm kiếm

        Returns:
            Danh sách tài liệu
        """
        return await self.document_repository.list(skip, limit, search)

    async def get_document(self, document_id: str) -> Tuple[ExcelDocumentInfo, bytes]:
        """
        Lấy thông tin và nội dung tài liệu.

        Args:
            document_id: ID của tài liệu

        Returns:
            Tuple chứa thông tin và nội dung tài liệu
        """
        return await self.document_repository.get(document_id)

    async def delete_document(self, document_id: str) -> None:
        """
        Xóa tài liệu.

        Args:
            document_id: ID của tài liệu
        """
        await self.document_repository.delete(document_id)

    async def convert_to_pdf(self, content: bytes, original_filename: str) -> Dict[str, Any]:
        """
        Chuyển đổi tài liệu Excel sang PDF.

        Args:
            content: Nội dung tài liệu Excel
            original_filename: Tên file gốc

        Returns:
            Dict chứa thông tin tài liệu PDF
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_excel:
                temp_excel.write(content)
                temp_excel_path = temp_excel.name

            pdf_filename = os.path.splitext(original_filename)[0] + ".pdf"
            temp_pdf_path = os.path.join(settings.TEMP_DIR, pdf_filename)

            try:
                import comtypes.client
                excel = comtypes.client.CreateObject('Excel.Application')
                excel.Visible = False
                excel.DisplayAlerts = False

                workbook = excel.Workbooks.Open(temp_excel_path)

                workbook.ExportAsFixedFormat(
                    Type=0,  
                    Filename=temp_pdf_path,
                    Quality=0,  
                    IncludeDocProperties=True,
                    IgnorePrintAreas=False
                )

                workbook.Close(False)
                excel.Quit()

                with open(temp_pdf_path, "rb") as f:
                    pdf_content = f.read()

                os.unlink(temp_excel_path)
                os.unlink(temp_pdf_path)

                document_repository = ExcelDocumentRepository(self.minio_client)
                document_info = ExcelDocumentInfo(
                    title=os.path.splitext(original_filename)[0],
                    description=f"PDF được chuyển đổi từ {original_filename}",
                    original_filename=pdf_filename,
                    file_size=len(pdf_content),
                    file_type="application/pdf",
                    storage_path="",  
                    metadata={"converted_from": original_filename}
                )

                document_info = await document_repository.save(document_info, pdf_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                if os.path.exists(temp_excel_path):
                    os.unlink(temp_excel_path)
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)

                raise ConversionException(f"Lỗi khi chuyển đổi sang PDF: {str(e)}")
        except Exception as e:
            raise ConversionException(str(e))

    async def convert_to_word(self, content: bytes, original_filename: str) -> Dict[str, Any]:
        """
        Chuyển đổi tài liệu Excel sang Word.

        Args:
            content: Nội dung tài liệu Excel
            original_filename: Tên file gốc

        Returns:
            Dict chứa thông tin tài liệu Word
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_excel:
                temp_excel.write(content)
                temp_excel_path = temp_excel.name

            word_filename = os.path.splitext(original_filename)[0] + ".docx"
            temp_word_path = os.path.join(settings.TEMP_DIR, word_filename)

            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.table import WD_TABLE_ALIGNMENT

                wb = load_workbook(temp_excel_path)
                sheet = wb.active

                doc = Document()

                doc.add_heading(os.path.splitext(original_filename)[0], 0)

                table = doc.add_table(rows=sheet.max_row, cols=sheet.max_column)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, row in enumerate(sheet.iter_rows()):
                    for j, cell in enumerate(row):
                        value = cell.value if cell.value is not None else ''
                        table.cell(i, j).text = str(value)

                doc.save(temp_word_path)

                with open(temp_word_path, "rb") as f:
                    word_content = f.read()

                os.unlink(temp_excel_path)
                os.unlink(temp_word_path)

                document_repository = ExcelDocumentRepository(self.minio_client)
                document_info = ExcelDocumentInfo(
                    title=os.path.splitext(original_filename)[0],
                    description=f"Word được chuyển đổi từ {original_filename}",
                    original_filename=word_filename,
                    file_size=len(word_content),
                    file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    storage_path="",  
                    metadata={"converted_from": original_filename}
                )

                document_info = await document_repository.save(document_info, word_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                if os.path.exists(temp_excel_path):
                    os.unlink(temp_excel_path)
                if os.path.exists(temp_word_path):
                    os.unlink(temp_word_path)

                raise ConversionException(f"Lỗi khi chuyển đổi sang Word: {str(e)}")
        except Exception as e:
            raise ConversionException(str(e))

    async def merge_documents(self, dto: MergeDocumentsDTO) -> Dict[str, Any]:
        """
        Gộp nhiều tài liệu Excel thành một.

        Args:
            dto: DTO chứa thông tin gộp tài liệu

        Returns:
            Dict chứa thông tin tài liệu đã gộp
        """
        try:
            merge_info = MergeInfo(
                id=str(uuid.uuid4()),
                document_ids=dto.document_ids,
                output_filename=dto.output_filename
            )

            await self.merge_repository.save(merge_info)

            documents = []
            for doc_id in dto.document_ids:
                try:
                    document_info, content = await self.document_repository.get(doc_id)
                    documents.append((document_info, content))
                except DocumentNotFoundException:
                    merge_info.status = "failed"
                    merge_info.error_message = f"Không tìm thấy tài liệu với ID: {doc_id}"
                    await self.merge_repository.update(merge_info)
                    raise

            merged_workbook = Workbook()
            merged_workbook.remove(merged_workbook.active)

            for document_info, content in documents:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
                    temp_file.write(content)
                    temp_file_path = temp_file.name

                try:
                    src_wb = load_workbook(temp_file_path)

                    for sheet_name in src_wb.sheetnames:
                        src_sheet = src_wb[sheet_name]

                        unique_sheet_name = f"{document_info.title}_{sheet_name}"[:31]  
                        if unique_sheet_name in merged_workbook.sheetnames:
                            unique_sheet_name = f"{unique_sheet_name}_{len(merged_workbook.sheetnames)}"[:31]

                        dest_sheet = merged_workbook.create_sheet(title=unique_sheet_name)

                        for row in src_sheet.iter_rows():
                            for cell in row:
                                dest_sheet.cell(
                                    row=cell.row,
                                    column=cell.column,
                                    value=cell.value
                                )

                                if cell.has_style:
                                    dest_cell = dest_sheet.cell(row=cell.row, column=cell.column)
                                    dest_cell.font = cell.font
                                    dest_cell.border = cell.border
                                    dest_cell.fill = cell.fill
                                    dest_cell.number_format = cell.number_format
                                    dest_cell.alignment = cell.alignment

                        for col_idx, col in enumerate(src_sheet.columns, 1):
                            col_letter = get_column_letter(col_idx)
                            if col_letter in src_sheet.column_dimensions:
                                dest_sheet.column_dimensions[col_letter].width = src_sheet.column_dimensions[
                                    col_letter].width

                        for row_idx, row in enumerate(src_sheet.rows, 1):
                            if row_idx in src_sheet.row_dimensions:
                                dest_sheet.row_dimensions[row_idx].height = src_sheet.row_dimensions[row_idx].height

                    src_wb.close()
                except Exception as e:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
                    raise MergeException(f"Lỗi khi gộp tài liệu: {str(e)}")

                os.unlink(temp_file_path)

            merged_file_path = os.path.join(settings.TEMP_DIR, dto.output_filename)
            merged_workbook.save(merged_file_path)

            with open(merged_file_path, "rb") as f:
                merged_content = f.read()

            os.unlink(merged_file_path)

            merged_document_info = ExcelDocumentInfo(
                title=os.path.splitext(dto.output_filename)[0],
                description=f"Tài liệu gộp từ {len(documents)} file Excel",
                original_filename=dto.output_filename,
                file_size=len(merged_content),
                file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                storage_path="",  
                metadata={"merged_from": [doc[0].id for doc in documents]}
            )

            merged_document_info = await self.document_repository.save(merged_document_info, merged_content)

            merge_info.status = "completed"
            merge_info.result_document_id = merged_document_info.id
            await self.merge_repository.update(merge_info)

            return {
                "id": merged_document_info.id,
                "filename": merged_document_info.original_filename,
                "file_size": merged_document_info.file_size
            }
        except Exception as e:
            merge_info.status = "failed"
            merge_info.error_message = str(e)
            await self.merge_repository.update(merge_info)
            raise MergeException(str(e))

    async def merge_documents_async(self, task_id: str, dto: MergeDocumentsDTO) -> None:
        """
        Gộp nhiều tài liệu Excel thành một (bất đồng bộ).

        Args:
            task_id: ID của tác vụ
            dto: DTO chứa thông tin gộp tài liệu
        """
        try:
            await self.rabbitmq_client.publish_merge_documents_task(
                task_id=task_id,
                document_ids=dto.document_ids,
                output_filename=dto.output_filename
            )
        except Exception as e:
            print(f"Lỗi khi đăng ký tác vụ gộp tài liệu: {str(e)}")

    async def get_merge_status(self, task_id: str) -> Dict[str, Any]:
        """
        Lấy trạng thái gộp tài liệu.

        Args:
            task_id: ID của tác vụ gộp tài liệu

        Returns:
            Dict chứa thông tin trạng thái
        """
        try:
            merge_info = await self.merge_repository.get(task_id)
            status_data = {
                "task_id": merge_info.id,
                "status": merge_info.status,
                "created_at": merge_info.created_at.isoformat(),
                "document_count": len(merge_info.document_ids),
                "output_filename": merge_info.output_filename
            }

            if merge_info.status == "completed" and merge_info.result_document_id:
                status_data["result_document_id"] = merge_info.result_document_id
                status_data["download_url"] = f"/documents/download/{merge_info.result_document_id}"
            elif merge_info.status == "failed" and merge_info.error_message:
                status_data["error_message"] = merge_info.error_message

            return status_data
        except Exception as e:
            return {
                "task_id": task_id,
                "status": "unknown",
                "error_message": str(e)
            }

    async def process_document_async(self, document_id: str) -> None:
        """
        Xử lý tài liệu bất đồng bộ.

        Args:
            document_id: ID của tài liệu
        """
        try:
            document_info, content = await self.document_repository.get(document_id)

            await self.rabbitmq_client.publish_convert_to_pdf_task(document_id)
        except Exception as e:
            print(f"Lỗi khi xử lý tài liệu {document_id}: {str(e)}")


class ExcelTemplateService:
    """
    Service xử lý mẫu tài liệu Excel.
    """

    def __init__(
            self,
            template_repository: ExcelTemplateRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            template_repository: Repository để làm việc với mẫu tài liệu
            minio_client: Client MinIO để lưu trữ mẫu tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.template_repository = template_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client
        self.batch_repository = BatchProcessingRepository()
        self.document_repository = ExcelDocumentRepository(minio_client)

    async def create_template(self, dto: CreateTemplateDTO, content: bytes) -> ExcelTemplateInfo:
        """
        Tạo mẫu tài liệu mới.

        Args:
            dto: DTO chứa thông tin mẫu tài liệu
            content: Nội dung mẫu tài liệu

        Returns:
            Thông tin mẫu tài liệu đã tạo
        """
        sheet_names = await self._get_sheet_names(content)

        template_info = ExcelTemplateInfo(
            name=dto.name,
            description=dto.description,
            category=dto.category,
            original_filename=dto.original_filename,
            file_size=len(content),
            storage_path="",  
            data_fields=dto.data_fields,
            metadata=dto.metadata,
            sheet_names=sheet_names
        )

        template_info = await self.template_repository.save(template_info, content)

        return template_info

    async def _get_sheet_names(self, content: bytes) -> List[str]:
        """
        Lấy danh sách tên sheet từ file Excel.

        Args:
            content: Nội dung file Excel

        Returns:
            Danh sách tên sheet
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                wb = load_workbook(temp_file_path, read_only=True)
                sheet_names = wb.sheetnames
                wb.close()

                os.unlink(temp_file_path)

                return sheet_names
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                print(f"Lỗi khi đọc tên sheet: {str(e)}")
                return []
        except Exception as e:
            print(f"Lỗi khi tạo file tạm: {str(e)}")
            return []

    async def get_templates(self, category: Optional[str] = None, skip: int = 0, limit: int = 10) -> List[
        ExcelTemplateInfo]:
        """
        Lấy danh sách mẫu tài liệu.

        Args:
            category: Danh mục để lọc
            skip: Số mẫu tài liệu bỏ qua
            limit: Số mẫu tài liệu tối đa trả về

        Returns:
            Danh sách mẫu tài liệu
        """
        return await self.template_repository.list(category, skip, limit)

    async def get_template(self, template_id: str) -> Tuple[ExcelTemplateInfo, bytes]:
        """
        Lấy thông tin và nội dung mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu

        Returns:
            Tuple chứa thông tin và nội dung mẫu tài liệu
        """
        return await self.template_repository.get(template_id)

    async def delete_template(self, template_id: str) -> None:
        """
        Xóa mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu
        """
        await self.template_repository.delete(template_id)

    async def apply_template(self, dto: TemplateDataDTO) -> Dict[str, Any]:
        """
        Áp dụng mẫu tài liệu với dữ liệu được cung cấp.

        Args:
            dto: DTO chứa thông tin để áp dụng vào mẫu

        Returns:
            Dict chứa thông tin tài liệu đã tạo
        """
        try:
            template_info, template_content = await self.template_repository.get(dto.template_id)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_template:
                temp_template.write(template_content)
                temp_template_path = temp_template.name

            result_filename = f"{template_info.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            temp_result_path = os.path.join(settings.TEMP_DIR, result_filename)

            try:
                wb = load_workbook(temp_template_path)

                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and "{{" in cell.value and "}}" in cell.value:
                                for key, value in dto.data.items():
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in cell.value:
                                        cell.value = cell.value.replace(placeholder, str(value))

                wb.save(temp_result_path)

                if dto.output_format.lower() == "pdf":
                    pdf_filename = os.path.splitext(result_filename)[0] + ".pdf"
                    temp_pdf_path = os.path.join(settings.TEMP_DIR, pdf_filename)

                    try:
                        import comtypes.client
                        excel = comtypes.client.CreateObject('Excel.Application')
                        excel.Visible = False
                        excel.DisplayAlerts = False

                        workbook = excel.Workbooks.Open(temp_result_path)

                        workbook.ExportAsFixedFormat(
                            Type=0,  
                            Filename=temp_pdf_path,
                            Quality=0,  
                            IncludeDocProperties=True,
                            IgnorePrintAreas=False
                        )

                        workbook.Close(False)
                        excel.Quit()

                        with open(temp_pdf_path, "rb") as f:
                            result_content = f.read()

                        os.unlink(temp_pdf_path)

                        result_filename = pdf_filename
                    except Exception as e:
                        raise TemplateApplicationException(f"Lỗi khi chuyển đổi sang PDF: {str(e)}")
                else:
                    with open(temp_result_path, "rb") as f:
                        result_content = f.read()

                os.unlink(temp_template_path)
                os.unlink(temp_result_path)

                document_info = ExcelDocumentInfo(
                    title=f"{template_info.name} - {datetime.now().strftime('%Y-%m-%d')}",
                    description=f"Tài liệu được tạo từ mẫu '{template_info.name}'",
                    original_filename=result_filename,
                    file_size=len(result_content),
                    file_type="application/pdf" if dto.output_format.lower() == "pdf" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    storage_path="",  
                    metadata={
                        "template_id": template_info.id,
                        "template_name": template_info.name,
                        "template_data": dto.data
                    }
                )

                document_info = await self.document_repository.save(document_info, result_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                if os.path.exists(temp_template_path):
                    os.unlink(temp_template_path)
                if os.path.exists(temp_result_path):
                    os.unlink(temp_result_path)

                raise TemplateApplicationException(f"Lỗi khi áp dụng mẫu: {str(e)}")
        except TemplateNotFoundException:
            raise
        except Exception as e:
            raise TemplateApplicationException(str(e))

    async def process_batch_async(self, task_id: str, template_id: str, content: bytes, filename: str,
                                  output_format: str) -> None:
        """
        Xử lý batch tài liệu.

        Args:
            task_id: ID của tác vụ
            template_id: ID của mẫu tài liệu
            content: Nội dung file dữ liệu (CSV, Excel)
            filename: Tên file dữ liệu
            output_format: Định dạng đầu ra
        """
        try:
            batch_info = BatchProcessingInfo(
                id=task_id,
                template_id=template_id,
                output_format=output_format
            )

            await self.batch_repository.save(batch_info)

            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                if filename.endswith('.csv'):
                    data_list = pd.read_csv(temp_file_path).to_dict('records')
                elif filename.endswith(('.xlsx', '.xls')):
                    data_list = pd.read_excel(temp_file_path).to_dict('records')
                else:
                    raise TemplateApplicationException(f"Định dạng file không được hỗ trợ: {filename}")

                batch_info.total_documents = len(data_list)
                await self.batch_repository.update(batch_info)

                result_documents = []

                for i, data in enumerate(data_list):
                    try:
                        template_data_dto = TemplateDataDTO(
                            template_id=template_id,
                            data=data,
                            output_format=output_format
                        )

                        result = await self.apply_template(template_data_dto)
                        result_documents.append(result)

                        batch_info.processed_documents = i + 1
                        await self.batch_repository.update(batch_info)
                    except Exception as e:
                        print(f"Lỗi khi xử lý bản ghi thứ {i}: {str(e)}")

                if output_format.lower() == "zip":
                    zip_filename = f"batch_{task_id}.zip"
                    temp_zip_path = os.path.join(settings.TEMP_DIR, zip_filename)

                    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for result in result_documents:
                            document_info, content = await self.document_repository.get(result["id"])
                            zipf.writestr(document_info.original_filename, content)

                    with open(temp_zip_path, "rb") as f:
                        zip_content = f.read()

                    zip_document_info = ExcelDocumentInfo(
                        title=f"Batch {task_id}",
                        description=f"File ZIP chứa {len(result_documents)} tài liệu được tạo từ mẫu",
                        original_filename=zip_filename,
                        file_size=len(zip_content),
                        file_type="application/zip",
                        storage_path="",  
                        metadata={
                            "template_id": template_id,
                            "batch_id": task_id,
                            "total_documents": len(result_documents)
                        }
                    )

                    zip_document_info = await self.document_repository.save(zip_document_info, zip_content)

                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    batch_info.result_file_id = zip_document_info.id
                    batch_info.result_file_path = zip_document_info.storage_path
                    await self.batch_repository.update(batch_info)

                    os.unlink(temp_zip_path)
                else:
                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    await self.batch_repository.update(batch_info)
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)

                raise TemplateApplicationException(f"Lỗi khi xử lý batch: {str(e)}")

            os.unlink(temp_file_path)
        except Exception as e:
            try:
                batch_info = await self.batch_repository.get(task_id)
                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)
            except:
                pass

            print(f"Lỗi khi xử lý batch {task_id}: {str(e)}")

    async def get_batch_status(self, batch_id: str) -> Dict[str, Any]:
        """
        Lấy trạng thái xử lý batch.

        Args:
            batch_id: ID của batch

        Returns:
            Dict chứa thông tin trạng thái
        """
        try:
            batch_info = await self.batch_repository.get(batch_id)
            status_data = {
                "task_id": batch_info.id,
                "status": batch_info.status,
                "created_at": batch_info.created_at.isoformat(),
                "total_documents": batch_info.total_documents,
                "processed_documents": batch_info.processed_documents,
                "output_format": batch_info.output_format
            }

            if batch_info.status == "completed" and batch_info.result_file_id:
                status_data["result_file_id"] = batch_info.result_file_id
                status_data["download_url"] = f"/documents/download/{batch_info.result_file_id}"
            elif batch_info.status == "failed" and batch_info.error_message:
                status_data["error_message"] = batch_info.error_message

            return status_data
        except Exception as e:
            return {
                "task_id": batch_id,
                "status": "unknown",
                "error_message": str(e)
            }