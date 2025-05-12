import os
import io
import tempfile
import asyncio
import uuid
import json
import pandas as pd
import zipfile
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

from application.dto import CreateDocumentDTO, CreateTemplateDTO, TemplateDataDTO, WatermarkDTO, BatchProcessingDTO
from domain.models import DocumentInfo, TemplateInfo, BatchProcessingInfo
from domain.exceptions import DocumentNotFoundException, TemplateNotFoundException, StorageException
from domain.exceptions import ConversionException, WatermarkException, TemplateApplicationException
from infrastructure.repository import DocumentRepository, TemplateRepository, BatchProcessingRepository
from infrastructure.minio_client import MinioClient
from infrastructure.rabbitmq_client import RabbitMQClient
from core.config import settings

# Thư viện xử lý Word
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import comtypes.client


class DocumentService:
    """
    Service xử lý tài liệu Word.
    """

    def __init__(
            self,
            document_repository: DocumentRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            document_repository: Repository để làm việc với tài liệu
            minio_client: Client MinIO để lưu trữ tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.document_repository = document_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client

    async def create_document(self, dto: CreateDocumentDTO, content: bytes) -> DocumentInfo:
        """
        Tạo tài liệu mới.

        Args:
            dto: DTO chứa thông tin tài liệu
            content: Nội dung tài liệu

        Returns:
            Thông tin tài liệu đã tạo
        """
        # Tạo thông tin tài liệu
        document_info = DocumentInfo(
            title=dto.title,
            description=dto.description,
            original_filename=dto.original_filename,
            file_size=len(content),
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if dto.original_filename.endswith(
                ".docx") else "application/msword",
            storage_path="",  # Sẽ được cập nhật sau khi lưu
            metadata=dto.metadata
        )

        # Lưu tài liệu
        document_info = await self.document_repository.save(document_info, content)

        return document_info

    async def get_documents(self, skip: int = 0, limit: int = 10, search: Optional[str] = None) -> List[DocumentInfo]:
        """
        Lấy danh sách tài liệu.

        Args:
            skip: Số tài liệu bỏ qua
            limit: Số tài liệu tối đa trả về
            search: Từ khóa tìm kiếm

        Returns:
            Danh sách tài liệu
        """
        return await self.document_repository.list(skip, limit, search)

    async def get_document(self, document_id: str) -> Tuple[DocumentInfo, bytes]:
        """
        Lấy thông tin và nội dung tài liệu.

        Args:
            document_id: ID của tài liệu

        Returns:
            Tuple chứa thông tin và nội dung tài liệu
        """
        return await self.document_repository.get(document_id)

    async def delete_document(self, document_id: str) -> None:
        """
        Xóa tài liệu.

        Args:
            document_id: ID của tài liệu
        """
        await self.document_repository.delete(document_id)

    async def convert_to_pdf(self, content: bytes, original_filename: str) -> Dict[str, Any]:
        """
        Chuyển đổi tài liệu Word sang PDF.

        Args:
            content: Nội dung tài liệu Word
            original_filename: Tên file gốc

        Returns:
            Dict chứa thông tin tài liệu PDF
        """
        try:
            # Tạo file tạm cho tài liệu Word
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
                temp_word.write(content)
                temp_word_path = temp_word.name

            # Tạo đường dẫn cho file PDF
            pdf_filename = os.path.splitext(original_filename)[0] + ".pdf"
            temp_pdf_path = os.path.join(settings.TEMP_DIR, pdf_filename)

            # Chuyển đổi Word sang PDF
            try:
                # Sử dụng COM để chuyển đổi
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = False

                # Mở file Word
                doc = word.Documents.Open(temp_word_path)

                # Lưu dưới dạng PDF
                doc.SaveAs(temp_pdf_path, FileFormat=17)  # 17 là mã định dạng PDF

                # Đóng file và ứng dụng Word
                doc.Close()
                word.Quit()

                # Đọc nội dung file PDF
                with open(temp_pdf_path, "rb") as f:
                    pdf_content = f.read()

                # Xóa file tạm
                os.unlink(temp_word_path)
                os.unlink(temp_pdf_path)

                # Tạo thông tin tài liệu PDF
                document_info = DocumentInfo(
                    title=os.path.splitext(original_filename)[0],
                    description=f"PDF được chuyển đổi từ {original_filename}",
                    original_filename=pdf_filename,
                    file_size=len(pdf_content),
                    file_type="application/pdf",
                    storage_path="",  # Sẽ được cập nhật sau khi lưu
                    metadata={"converted_from": original_filename}
                )

                # Lưu tài liệu PDF
                document_info = await self.document_repository.save(document_info, pdf_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                # Xóa file tạm nếu có lỗi
                if os.path.exists(temp_word_path):
                    os.unlink(temp_word_path)
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)

                raise ConversionException(f"Lỗi khi chuyển đổi sang PDF: {str(e)}")
        except Exception as e:
            raise ConversionException(str(e))

    async def add_watermark(self, content: bytes, original_filename: str, dto: WatermarkDTO) -> Dict[str, Any]:
        """
        Thêm watermark vào tài liệu Word.

        Args:
            content: Nội dung tài liệu Word
            original_filename: Tên file gốc
            dto: DTO chứa thông tin watermark

        Returns:
            Dict chứa thông tin tài liệu đã thêm watermark
        """
        try:
            # Tạo file tạm cho tài liệu Word
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
                temp_word.write(content)
                temp_word_path = temp_word.name

            # Tạo đường dẫn cho file kết quả
            watermark_filename = os.path.splitext(original_filename)[0] + "_watermark.docx"
            temp_result_path = os.path.join(settings.TEMP_DIR, watermark_filename)

            try:
                # Mở tài liệu Word
                doc = Document(temp_word_path)

                # Thêm watermark vào tài liệu
                for section in doc.sections:
                    header = section.header
                    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

                    # Xóa nội dung cũ
                    paragraph.clear()

                    # Thêm watermark
                    run = paragraph.add_run(dto.text)
                    run.font.size = Pt(40)
                    run.font.color.rgb = RGBColor(192, 192, 192)  # Màu xám
                    run.font.bold = True

                    # Thiết lập căn giữa
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Lưu tài liệu
                doc.save(temp_result_path)

                # Đọc nội dung file kết quả
                with open(temp_result_path, "rb") as f:
                    result_content = f.read()

                # Xóa file tạm
                os.unlink(temp_word_path)
                os.unlink(temp_result_path)

                # Tạo thông tin tài liệu kết quả
                document_info = DocumentInfo(
                    title=os.path.splitext(original_filename)[0] + " (Watermark)",
                    description=f"Tài liệu với watermark '{dto.text}'",
                    original_filename=watermark_filename,
                    file_size=len(result_content),
                    file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    storage_path="",  # Sẽ được cập nhật sau khi lưu
                    metadata={"watermark": dto.text, "original_filename": original_filename}
                )

                # Lưu tài liệu kết quả
                document_info = await self.document_repository.save(document_info, result_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                # Xóa file tạm nếu có lỗi
                if os.path.exists(temp_word_path):
                    os.unlink(temp_word_path)
                if os.path.exists(temp_result_path):
                    os.unlink(temp_result_path)

                raise WatermarkException(f"Lỗi khi thêm watermark: {str(e)}")
        except Exception as e:
            raise WatermarkException(str(e))

    async def process_document_async(self, document_id: str) -> None:
        """
        Xử lý tài liệu bất đồng bộ.

        Args:
            document_id: ID của tài liệu
        """
        try:
            # Lấy thông tin tài liệu
            document_info, content = await self.document_repository.get(document_id)

            # Gửi tin nhắn để xử lý tài liệu
            await self.rabbitmq_client.publish_convert_to_pdf_task(document_id)
        except Exception as e:
            # Log lỗi
            print(f"Lỗi khi xử lý tài liệu {document_id}: {str(e)}")


class TemplateService:
    """
    Service xử lý mẫu tài liệu Word.
    """

    def __init__(
            self,
            template_repository: TemplateRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            template_repository: Repository để làm việc với mẫu tài liệu
            minio_client: Client MinIO để lưu trữ mẫu tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.template_repository = template_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client
        self.batch_repository = BatchProcessingRepository()

    async def create_template(self, dto: CreateTemplateDTO, content: bytes) -> TemplateInfo:
        """
        Tạo mẫu tài liệu mới.

        Args:
            dto: DTO chứa thông tin mẫu tài liệu
            content: Nội dung mẫu tài liệu

        Returns:
            Thông tin mẫu tài liệu đã tạo
        """
        # Tạo thông tin mẫu tài liệu
        template_info = TemplateInfo(
            name=dto.name,
            description=dto.description,
            category=dto.category,
            original_filename=dto.original_filename,
            file_size=len(content),
            storage_path="",  # Sẽ được cập nhật sau khi lưu
            data_fields=dto.data_fields,
            metadata=dto.metadata
        )

        # Lưu mẫu tài liệu
        template_info = await self.template_repository.save(template_info, content)

        return template_info

    async def get_templates(self, category: Optional[str] = None, skip: int = 0, limit: int = 10) -> List[TemplateInfo]:
        """
        Lấy danh sách mẫu tài liệu.

        Args:
            category: Danh mục để lọc
            skip: Số mẫu tài liệu bỏ qua
            limit: Số mẫu tài liệu tối đa trả về

        Returns:
            Danh sách mẫu tài liệu
        """
        return await self.template_repository.list(category, skip, limit)

    async def get_template(self, template_id: str) -> Tuple[TemplateInfo, bytes]:
        """
        Lấy thông tin và nội dung mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu

        Returns:
            Tuple chứa thông tin và nội dung mẫu tài liệu
        """
        return await self.template_repository.get(template_id)

    async def delete_template(self, template_id: str) -> None:
        """
        Xóa mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu
        """
        await self.template_repository.delete(template_id)

    async def apply_template(self, dto: TemplateDataDTO) -> Dict[str, Any]:
        """
        Áp dụng mẫu tài liệu với dữ liệu được cung cấp.

        Args:
            dto: DTO chứa thông tin để áp dụng vào mẫu

        Returns:
            Dict chứa thông tin tài liệu đã tạo
        """
        try:
            # Lấy thông tin và nội dung mẫu tài liệu
            template_info, template_content = await self.template_repository.get(dto.template_id)

            # Tạo file tạm cho mẫu tài liệu
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_template:
                temp_template.write(template_content)
                temp_template_path = temp_template.name

            # Tạo đường dẫn cho file kết quả
            result_filename = f"{template_info.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            temp_result_path = os.path.join(settings.TEMP_DIR, result_filename)

            try:
                # Mở mẫu tài liệu
                doc = Document(temp_template_path)

                # Áp dụng dữ liệu vào mẫu
                for paragraph in doc.paragraphs:
                    for key, value in dto.data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

                # Áp dụng dữ liệu vào các bảng
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in dto.data.items():
                                    if f"{{{{{key}}}}}" in paragraph.text:
                                        paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

                # Lưu tài liệu
                doc.save(temp_result_path)

                # Xử lý output format
                if dto.output_format.lower() == "pdf":
                    # Chuyển đổi sang PDF
                    pdf_filename = os.path.splitext(result_filename)[0] + ".pdf"
                    temp_pdf_path = os.path.join(settings.TEMP_DIR, pdf_filename)

                    # Sử dụng COM để chuyển đổi
                    word_app = comtypes.client.CreateObject('Word.Application')
                    word_app.Visible = False

                    # Mở file Word
                    word_doc = word_app.Documents.Open(temp_result_path)

                    # Lưu dưới dạng PDF
                    word_doc.SaveAs(temp_pdf_path, FileFormat=17)  # 17 là mã định dạng PDF

                    # Đóng file và ứng dụng Word
                    word_doc.Close()
                    word_app.Quit()

                    # Đọc nội dung file PDF
                    with open(temp_pdf_path, "rb") as f:
                        result_content = f.read()

                    # Xóa file tạm
                    os.unlink(temp_pdf_path)

                    # Cập nhật tên file
                    result_filename = pdf_filename
                else:
                    # Đọc nội dung file Word
                    with open(temp_result_path, "rb") as f:
                        result_content = f.read()

                # Xóa file tạm
                os.unlink(temp_template_path)
                os.unlink(temp_result_path)

                # Tạo thông tin tài liệu kết quả
                document_info = DocumentInfo(
                    title=f"{template_info.name} - {datetime.now().strftime('%Y-%m-%d')}",
                    description=f"Tài liệu được tạo từ mẫu '{template_info.name}'",
                    original_filename=result_filename,
                    file_size=len(result_content),
                    file_type="application/pdf" if dto.output_format.lower() == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    storage_path="",  # Sẽ được cập nhật sau khi lưu
                    metadata={
                        "template_id": template_info.id,
                        "template_name": template_info.name,
                        "template_data": dto.data
                    }
                )

                # Lưu tài liệu kết quả (sử dụng Document Repository)
                document_repository = DocumentRepository(self.minio_client)
                document_info = await document_repository.save(document_info, result_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                # Xóa file tạm nếu có lỗi
                if os.path.exists(temp_template_path):
                    os.unlink(temp_template_path)
                if os.path.exists(temp_result_path):
                    os.unlink(temp_result_path)

                raise TemplateApplicationException(f"Lỗi khi áp dụng mẫu: {str(e)}")
        except TemplateNotFoundException:
            raise
        except Exception as e:
            raise TemplateApplicationException(str(e))

    async def process_batch_async(self, task_id: str, template_id: str, content: bytes, filename: str,
                                  output_format: str) -> None:
        """
        Xử lý batch tài liệu.

        Args:
            task_id: ID của tác vụ
            template_id: ID của mẫu tài liệu
            content: Nội dung file dữ liệu (CSV, Excel)
            filename: Tên file dữ liệu
            output_format: Định dạng đầu ra
        """
        try:
            # Lưu thông tin batch
            batch_info = BatchProcessingInfo(
                id=task_id,
                template_id=template_id,
                output_format=output_format
            )

            await self.batch_repository.save(batch_info)

            # Đọc dữ liệu từ file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                # Đọc dữ liệu
                if filename.endswith('.csv'):
                    data_list = pd.read_csv(temp_file_path).to_dict('records')
                elif filename.endswith(('.xlsx', '.xls')):
                    data_list = pd.read_excel(temp_file_path).to_dict('records')
                else:
                    raise TemplateApplicationException(f"Định dạng file không được hỗ trợ: {filename}")

                # Cập nhật thông tin batch
                batch_info.total_documents = len(data_list)
                await self.batch_repository.update(batch_info)

                # Tạo một tài liệu cho mỗi bản ghi
                result_documents = []

                # Xử lý từng bản ghi
                for i, data in enumerate(data_list):
                    try:
                        # Áp dụng mẫu
                        template_data_dto = TemplateDataDTO(
                            template_id=template_id,
                            data=data,
                            output_format=output_format
                        )

                        result = await self.apply_template(template_data_dto)
                        result_documents.append(result)

                        # Cập nhật tiến trình
                        batch_info.processed_documents = i + 1
                        await self.batch_repository.update(batch_info)
                    except Exception as e:
                        print(f"Lỗi khi xử lý bản ghi thứ {i}: {str(e)}")

                # Nếu output_format là zip, tạo file zip chứa tất cả các tài liệu
                if output_format.lower() == "zip":
                    document_repository = DocumentRepository(self.minio_client)

                    # Tạo file zip
                    zip_filename = f"batch_{task_id}.zip"
                    temp_zip_path = os.path.join(settings.TEMP_DIR, zip_filename)

                    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for result in result_documents:
                            document_info, content = await document_repository.get(result["id"])
                            zipf.writestr(document_info.original_filename, content)

                    # Đọc nội dung file zip
                    with open(temp_zip_path, "rb") as f:
                        zip_content = f.read()

                    # Tạo thông tin tài liệu zip
                    zip_document_info = DocumentInfo(
                        title=f"Batch {task_id}",
                        description=f"File ZIP chứa {len(result_documents)} tài liệu được tạo từ mẫu",
                        original_filename=zip_filename,
                        file_size=len(zip_content),
                        file_type="application/zip",
                        storage_path="",  # Sẽ được cập nhật sau khi lưu
                        metadata={
                            "template_id": template_id,
                            "batch_id": task_id,
                            "total_documents": len(result_documents)
                        }
                    )

                    # Lưu tài liệu zip
                    zip_document_info = await document_repository.save(zip_document_info, zip_content)

                    # Cập nhật thông tin batch
                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    batch_info.result_file_id = zip_document_info.id
                    batch_info.result_file_path = zip_document_info.storage_path
                    await self.batch_repository.update(batch_info)

                    # Xóa file tạm
                    os.unlink(temp_zip_path)
                else:
                    # Cập nhật thông tin batch
                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    await self.batch_repository.update(batch_info)
            except Exception as e:
                # Xóa file tạm
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

                # Cập nhật thông tin batch
                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)

                raise TemplateApplicationException(f"Lỗi khi xử lý batch: {str(e)}")

            # Xóa file tạm
            os.unlink(temp_file_path)
        except Exception as e:
            # Lưu thông tin lỗi
            try:
                batch_info = await self.batch_repository.get(task_id)
                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)
            except:
                pass

            print(f"Lỗi khi xử lý batch {task_id}: {str(e)}")

    async def get_batch_status(self, batch_id: str) -> BatchProcessingInfo:
        """
        Lấy trạng thái xử lý batch.

        Args:
            batch_id: ID của batch

        Returns:
            Thông tin trạng thái batch
        """
        return await self.batch_repository.get(batch_id)